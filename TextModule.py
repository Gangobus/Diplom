# TextModule.py
import os
import sys
if sys.__stdout__ is None or sys.__stderr__ is None:
    os.environ['KIVY_NO_CONSOLELOG'] = '1'
from docx import Document
from tkinter import Tk, filedialog
import keyboard
import subprocess
import time
from kivy.clock import Clock
from kivy.uix.boxlayout import BoxLayout
from zipfile import BadZipFile
import win32com.client as win32

#основной класс для работы с текстом
class Text1Class():
    def __init__(self, shortcut_labels=None, **kwargs):
        super().__init__(**kwargs)
        self.secondtext = Text2Class()
        self.formatted_time = "00:00:00"
        if shortcut_labels is None:
            shortcut_labels = {
                'Alt+`': 'М? — ',
                'Alt+ё': 'М? — ',
                'Alt+1': 'M1 — ',
                'Alt+2': 'M2 — ',
                'Alt+3': 'M3 — ',
                'Alt+4': 'M4 — ',
                'Alt+5': '<Речь неразборчива>',
                'Alt+6': 'Н — ',
                'Alt+7': 'О — ',
                'Alt+8': '<>',
                'Alt+9': 'Спорная фонограмма №',

                'Ctrl+`': 'Ж? — ',
                'Ctrl+ё': 'Ж? — ',
                'Ctrl+1': 'Ж1 — ',
                'Ctrl+2': 'Ж2 — ',
                'Ctrl+3': 'Ж3 — ',
                'Ctrl+4': 'Ж4 — ',
                'Ctrl+5': 'Р — ',
                'Ctrl+6': '<>',
                'Ctrl+7': '<>',
                'Ctrl+8': '<>',
                'Ctrl+9': '<>',
                'Ctrl+0': '<Комментарий эксперта>'
            }
        self.shortcut_labels = shortcut_labels
        try:
            keyboard.add_hotkey("Ctrl+ё", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+ё'],))
            keyboard.add_hotkey("Ctrl+1", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+1'],))
            keyboard.add_hotkey("Ctrl+2", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+2'],))
            keyboard.add_hotkey("Ctrl+3", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+3'],))
            keyboard.add_hotkey("Ctrl+4", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+4'],))
            keyboard.add_hotkey("Ctrl+5", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+5'],))
            keyboard.add_hotkey("Ctrl+6", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+6'],))
            keyboard.add_hotkey("Ctrl+7", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+7'],))
            keyboard.add_hotkey("Ctrl+8", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+8'],))
            keyboard.add_hotkey("Ctrl+9", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+9'],))
            keyboard.add_hotkey("Ctrl+0", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+0'],))
            keyboard.add_hotkey("Alt+ё", self.insert_text_from_shortcut, args=(self.shortcut_labels['Alt+ё'],))
            keyboard.add_hotkey("Alt+1", self.insert_text_from_shortcut, args=(self.shortcut_labels['Alt+1'],))
            keyboard.add_hotkey("Alt+2", self.insert_text_from_shortcut, args=(self.shortcut_labels['Alt+2'],))
            keyboard.add_hotkey("Alt+3", self.insert_text_from_shortcut, args=(self.shortcut_labels['Alt+3'],))
            keyboard.add_hotkey("Alt+4", self.insert_text_from_shortcut, args=(self.shortcut_labels['Alt+4'],))
            keyboard.add_hotkey("Alt+5", self.insert_text_from_shortcut, args=(self.shortcut_labels['Alt+5'],))
            keyboard.add_hotkey("Alt+6", self.insert_text_from_shortcut, args=(self.shortcut_labels['Alt+6'],))
            keyboard.add_hotkey("Alt+7", self.insert_text_from_shortcut, args=(self.shortcut_labels['Alt+7'],))
            keyboard.add_hotkey("Alt+8", self.insert_text_from_shortcut, args=(self.shortcut_labels['Alt+8'],))
            keyboard.add_hotkey("Alt+9", self.insert_text_from_shortcut, args=(self.shortcut_labels['Alt+9'],))
            keyboard.add_hotkey("Alt+0", self.print_formatted_time)
        except:
            keyboard.add_hotkey("Ctrl+`", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+`'],))
            keyboard.add_hotkey("Ctrl+1", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+1'],))
            keyboard.add_hotkey("Ctrl+2", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+2'],))
            keyboard.add_hotkey("Ctrl+3", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+3'],))
            keyboard.add_hotkey("Ctrl+4", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+4'],))
            keyboard.add_hotkey("Ctrl+5", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+5'],))
            keyboard.add_hotkey("Ctrl+6", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+6'],))
            keyboard.add_hotkey("Ctrl+7", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+7'],))
            keyboard.add_hotkey("Ctrl+8", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+8'],))
            keyboard.add_hotkey("Ctrl+9", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+9'],))
            keyboard.add_hotkey("Ctrl+0", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+0'],))
            keyboard.add_hotkey("Alt+`", self.insert_text_from_shortcut, args=(self.shortcut_labels['Alt+`'],))
            keyboard.add_hotkey("Alt+1", self.insert_text_from_shortcut, args=(self.shortcut_labels['Alt+1'],))
            keyboard.add_hotkey("Alt+2", self.insert_text_from_shortcut, args=(self.shortcut_labels['Alt+2'],))
            keyboard.add_hotkey("Alt+3", self.insert_text_from_shortcut, args=(self.shortcut_labels['Alt+3'],))
            keyboard.add_hotkey("Alt+4", self.insert_text_from_shortcut, args=(self.shortcut_labels['Alt+4'],))
            keyboard.add_hotkey("Alt+5", self.insert_text_from_shortcut, args=(self.shortcut_labels['Alt+5'],))
            keyboard.add_hotkey("Alt+6", self.insert_text_from_shortcut, args=(self.shortcut_labels['Alt+6'],))
            keyboard.add_hotkey("Alt+7", self.insert_text_from_shortcut, args=(self.shortcut_labels['Alt+7'],))
            keyboard.add_hotkey("Alt+8", self.insert_text_from_shortcut, args=(self.shortcut_labels['Alt+8'],))
            keyboard.add_hotkey("Alt+9", self.insert_text_from_shortcut, args=(self.shortcut_labels['Alt+9'],))
            keyboard.add_hotkey("Alt+0", self.print_formatted_time)

    #вствка текста зависящего от нажатой кнопки в основной текст
    def add_text_to_input(self):
        a1 = self.ids.txt1.text
        cursor_pos = self.ids.txt1.cursor[0]
        line_start = a1.rfind('\n', 0, cursor_pos) + 1
        text_to_insert = ""

        def insert_text(dt):
            self.ids.txt1.insert_text(text_to_insert)
        Clock.schedule_once(insert_text)

    #метод для обновления текушего времени воспроизведения
    def set_formatted_time(self, formatted_time):
        self.formatted_time = formatted_time

    # метод для вставки текушего времени воспроизведения в качестве текстовой метки
    def print_formatted_time(self):
        time_text = ("<" + self.formatted_time + ">")
        Clock.schedule_once(lambda dt: self.ids.txt1.insert_text(time_text), 0)

    #метод для вставки текстовых меток
    def insert_text_from_shortcut(self, text_to_insert):
        Clock.schedule_once(lambda dt: self.ids.txt1.insert_text(text_to_insert), 0)

    #метод для перревода интерфейса в режим работы с двумя текстами
    def double_text_mode(self):
        secondtext = Text2Class()
        if self.secondtext in self.ids.bta.children:
            self.ids.bta.remove_widget(self.secondtext)
        else:
            self.ids.bta.add_widget(self.secondtext)
    def on_open_txt_file_dialog(self, *args):
        Clock.schedule_once(self.open_txt_file_dialog)

    #открытие текстовых файлов через диалоговое окно
    def open_txt_file_dialog(self, *args):
        try:
            root = Tk()
            root.withdraw()
            txt1_o_file_path = filedialog.askopenfilename(filetypes=[("Text Files", "*.txt"),
                                                              ("Word Documents", "*.docx"),
                                                              ("Word Documents (Legacy)", "*.doc"),
                                                              ('All Files', '*.*')])
            if txt1_o_file_path:
                if txt1_o_file_path.endswith(".docx"):
                    document_1_o = Document(txt1_o_file_path)
                    paragraphs = [para.text for para in document_1_o.paragraphs]
                    text = "\n".join(paragraphs)
                    self.ids.txt1.text = text
                elif txt1_o_file_path.endswith(".doc"):
                    word = win32.gencache.EnsureDispatch("Word.Application")
                    word.Visible = False
                    doc = word.Documents.Open(txt1_o_file_path)
                    text = doc.Content.Text
                    self.ids.txt1.text = text
                    doc.Close()
                    word.Quit()
                else:
                    with open(txt1_o_file_path, 'r', encoding='utf-8') as file:
                        text = file.read()
                        self.ids.txt1.text = text
        except BadZipFile:
            pass
        except FileNotFoundError:
            print(FileNotFoundError)
            pass
        except IOError:
            print(IOError)
            pass
        except PermissionError:
            print(PermissionError)
            pass
        except Exception as e:
            print(e)
            pass
    def on_save_text_to_file(self, *args):
        Clock.schedule_once(self.save_text_to_file)
    #сохранение текстовых файлов через диалоговое окно
    def save_text_to_file(self, *args):
        try:
            txt_1_s_file_path = filedialog.asksaveasfilename(defaultextension=".txt",
                                                     filetypes=[("Text Files", "*.txt"),
                                                                ("Word Documents", "*.docx"),
                                                                ("Word Documents (Legacy)", "*.doc")])
            if txt_1_s_file_path:
                if txt_1_s_file_path.endswith(".docx"):
                    document = Document()
                    document.add_paragraph(self.ids.txt1.text)
                    document.save(txt_1_s_file_path)
                elif txt_1_s_file_path.endswith(".doc"):
                    word = win32.gencache.EnsureDispatch("Word.Application")
                    word.Visible = False
                    doc = word.Documents.Add()
                    doc.Content.Text = self.ids.txt1.text
                    doc.SaveAs(txt_1_s_file_path)
                    doc.Close()
                    word.Quit()
                else:
                    with open(txt_1_s_file_path, 'w', encoding='utf-8') as file:
                        file.write(self.ids.txt1.text)
        except BadZipFile:
            print(BadZipFile)
            pass
        except FileNotFoundError:
            print(FileNotFoundError)
            pass
        except IOError:
            print(IOError)
            pass
        except PermissionError:
            print(PermissionError)
            pass
        except Exception as e:
            print(e)
            pass

    #увеличение размера шрифта первого текста
    def increase_font_size1(self):
        self.ids.txt1.font_size += 2

    #уменьшение размера шрифта первого текста
    def decrease_font_size1(self):
        self.ids.txt1.font_size -= 2

    # запуск модуля распознования речи
    def spech_recogn(self):
        current_dir = os.path.dirname(os.path.abspath(__file__))
        exe_dir = os.path.join(current_dir, "auto_text_from_audio")
        exe_path = os.path.join(exe_dir, "voice-recognition.exe")
        files_dir = os.path.join(exe_dir, "vosk-model-ru-0.42")
        subprocess.Popen([exe_path, files_dir], creationflags=subprocess.CREATE_NEW_CONSOLE, cwd=exe_dir)
        # subprocess.Popen([exe_path, files_dir], cwd=exe_dir)

    def txt2(self):
        secondtext = Text2Class()
        if self.secondtext in self.ids.bta.children:
            self.ids.bta.remove_widget(self.secondtext)
        else:
            self.ids.bta.add_widget(self.secondtext)

#класс для работы с вторым текстом
class Text2Class(BoxLayout):
    def on_open_2_file_dialog(self, *args):
        Clock.schedule_once(self.open_2_file_dialog)
    #открытие текстового файла второго текста
    def open_2_file_dialog(self, *args):
        try:
            root = Tk()
            root.withdraw()
            txt_2_o_file_path = filedialog.askopenfilename(filetypes=[("Text Files", "*.txt"),
                                                              ("Word Documents", "*.docx"),
                                                              ("Word Documents (Legacy)", "*.doc"),
                                                              ('All Files', '*.*')])
            if txt_2_o_file_path:
                if txt_2_o_file_path.endswith(".docx"):
                    document = Document(txt_2_o_file_path)
                    paragraphs = [para.text for para in document.paragraphs]
                    text = "\n".join(paragraphs)
                    self.ids.txt22.text = text
                elif txt_2_o_file_path.endswith(".doc"):
                    word = win32.gencache.EnsureDispatch("Word.Application")
                    word.Visible = False
                    doc = word.Documents.Open(txt_2_o_file_path)
                    text = doc.Content.Text
                    self.ids.txt22.text = text
                    doc.Close()
                    word.Quit()
                else:
                    with open(txt_2_o_file_path, 'r', encoding='utf-8') as file:
                        text = file.read()
                        self.ids.txt22.text = text
        except BadZipFile:
            pass
        except FileNotFoundError:
            print(FileNotFoundError)
            pass
        except IOError:
            print(IOError)
            pass
        except PermissionError:
            print(PermissionError)
            pass
        except Exception as e:
            print(e)
            pass

    def on_save_2_text_to_file(self, *args):
        Clock.schedule_once(self.save_2_text_to_file)
    # сохранение текстового файла для второго текста
    def save_2_text_to_file(self, *args):
        try:
            txt_2_s_file_path = filedialog.asksaveasfilename(defaultextension=".txt",
                                                     filetypes=[("Text Files", "*.txt"),
                                                                ("Word Documents", "*.docx"),
                                                                ("Word Documents (Legacy)", "*.doc")])
            if txt_2_s_file_path:
                if txt_2_s_file_path.endswith(".docx"):
                    document = Document()
                    document.add_paragraph(self.ids.txt22.text)
                    document.save(txt_2_s_file_path)
                elif txt_2_s_file_path.endswith(".doc"):
                    word = win32.gencache.EnsureDispatch("Word.Application")
                    word.Visible = False
                    doc = word.Documents.Add()
                    doc.Content.Text = self.ids.txt22.text
                    doc.SaveAs(txt_2_s_file_path)
                    doc.Close()
                    word.Quit()
                else:
                    with open(txt_2_s_file_path, 'w', encoding='utf-8') as file:
                        file.write(self.ids.txt22.text)
        except BadZipFile:
            pass
        except FileNotFoundError:
            print(FileNotFoundError)
            pass
        except IOError:
            print(IOError)
            pass
        except PermissionError:
            print(PermissionError)
            pass
        except Exception as e:
            print(e)
            pass
    #увеличение размера шрифта второго текста
    def increase_font_size2(self):
        self.ids.txt22.font_size += 2

    #уменьшение размера шрифта второго текста
    def decrease_font_size2(self):
        self.ids.txt22.font_size -= 2
